from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io


def build_cover_letter_docx(cover_letter_text: str, job_title: str, company: str) -> bytes:
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f"Cover Letter — {job_title} at {company}")
    run.bold = True
    run.font.size = Pt(13)

    doc.add_paragraph()

    # Body
    for para_text in cover_letter_text.split("\n\n"):
        para_text = para_text.strip()
        if not para_text:
            continue
        p = doc.add_paragraph(para_text)
        p.paragraph_format.space_after  = Pt(8)
        p.paragraph_format.line_spacing = Pt(14)
        for run in p.runs:
            run.font.size = Pt(11)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()